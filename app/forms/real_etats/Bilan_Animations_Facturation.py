# coding: utf-8

# Imports
from django import forms

class Bilan_Animations_Facturation(forms.Form):

	# Import
	from smmaranim.custom_settings import EMPTY_VALUE

	# Filtres

	zl_presta = forms.ChoiceField(
		choices=[EMPTY_VALUE],
		label='Organisme'
	)

	zl_marche_lot = forms.ChoiceField(
		choices=[EMPTY_VALUE], label='Lot'
	)

	zd_date_debut = forms.DateField(label='Animations du')

	zd_date_fin = forms.DateField(label='Animations au')

	# Méthodes Django

	def __init__(self, *args, **kwargs):

		# Imports
		from app.models import TOrganisme
		from app.models import TPrestatairesMarche
		from app.models import TUtilisateur

		# Arguments
		self.rq = kwargs.pop('kwarg_rq')

		super().__init__(*args, **kwargs)

		# Organismes sélectionnables
		self.fields['zl_presta'].choices += [
			(oOrg.pk, oOrg) for oOrg in TOrganisme.objects.all()
		]

		# Lots sélectionnables (tous si utilisateur SMMAR ou ceux
		# concernant l'organisme de l'utilisateur connecté)
		oOrg = TUtilisateur.get_util_connect(self.rq).get_org()
		if oOrg.get_est_prest():
			qsPm = oOrg.get_pm()
		else:
			qsPm = TPrestatairesMarche.objects.all()
		marcheLots = []
		for oPm in qsPm:
			marcheLots.append((oPm.pk, oPm))
		self.fields['zl_marche_lot'].choices += marcheLots

	# Méthodes privées

	def __cleaned_data(self):

		"""Récupération des données nettoyées du formulaire"""

		# Imports
		from app.models import TOrganisme
		from datetime import datetime

		# Clés
		keys = [
			'zl_presta',
			'zl_marche_lot',
			'zd_date_debut',
			'zd_date_fin'
		]

		# Édition bilan
		if self.rq.GET.get('action') == 'editer-bilan':
			return {
				'zl_presta': TOrganisme.objects.get(
					pk=self.rq.GET.get('presta')
				),
				'zl_marche_lot': self.rq.GET.get('marcheLot'),
				'zd_date_debut': datetime.strptime(
					self.rq.GET.get('marcheLotBilanPeriodeDu'),
					'%Y-%m-%d'
				).date(),
				'zd_date_fin': datetime.strptime(
					self.rq.GET.get('marcheLotBilanPeriodeAu'),
					'%Y-%m-%d'
				).date()
			}

		# Formulaire HTML
		else:
			if not self.data:
				return {
					element: self.fields[element].initial for element in keys
				}
			else:
				return {
					element: self.cleaned_data.get(element) for element in keys
				}

	def __edit_bilan(self):

		"""Bilan des animations au format Word"""

		# Imports
		from app.functions.gen_string import sub as gen_string
		from app.functions.get_local_format import sub as get_local_format
		from app.models import TPrestatairesMarche
		from app.models import TProjet
		from django.http import HttpResponse
		from docx import Document
		from docx.enum.section import WD_SECTION
		from docx.enum.text import WD_ALIGN_PARAGRAPH
		from docx.shared import Mm
		from smmaranim.custom_settings import TEMPLATES_ROOT

		def insert_paragraph_after(paragraph):

			"""Insertion d'un nouveau paragraphe après un paragraphe"""

			# Imports
			from docx.text.paragraph import Paragraph
			from docx.oxml.xmlchemy import OxmlElement

			newP = OxmlElement('w:p')
			paragraph._p.addnext(newP)
			newParagraph = Paragraph(newP, paragraph._parent)

			return newParagraph

		# Ouverture du modèle Word
		document = Document(TEMPLATES_ROOT + '/Animations_Bilan.docx')

		# Récupération des données nettoyées du formulaire
		cleaned_data = self.__cleaned_data()
		presta = cleaned_data['zl_presta']
		marcheLot = cleaned_data['zl_marche_lot']
		marcheLotBilanPeriodeDu = cleaned_data['zd_date_debut']
		marcheLotBilanPeriodeAu = cleaned_data['zd_date_fin']

		# Instance TPrestatairesMarche
		oPm = TPrestatairesMarche.objects.get(pk=marcheLot)

		# Données fixes à automatiser
		replace = {
			'MarcheNumero': oPm.get_marche().get_int_marche(),
			'MarcheLotNumero': oPm.get_numero_lot(),
			'MarcheLotBilanPeriode': 'Du {} au {}'.format(
				get_local_format(marcheLotBilanPeriodeDu),
				get_local_format(marcheLotBilanPeriodeAu)
			),
			'MarcheLotPrestataire': presta
		}

		# Pour chaque paragraphe du modèle Word, automatisation des
		# données fixes (rechercher-remplacer)
		for paragraph in document.paragraphs:
			for k, v in replace.items():
				valueToReplace = '${}$'.format(k)
				if valueToReplace in paragraph.text:
					paragraph.text = paragraph.text.replace(
						valueToReplace, str(v)
					)

		#bodySection = document.sections[0]
		#bodySection.start_type = WD_SECTION.NEW_PAGE

		# Animations
		qsAni = self.__get_queryset()

		# Projets
		qsPro = TProjet.objects.filter(
			pk__in=qsAni.values_list('id_projet', flat=True)
		)

		# Récupération du paragraphe $START$
		for paragraph in document.paragraphs:
			if paragraph.text == '$START$':
				documentStart = paragraph
				paragraphToRemove = paragraph

		# Pour chaque projet...
		for oPro in qsPro:

			# Projet
			paraTitre1 = insert_paragraph_after(documentStart)
			paraTitre1.style = document.styles['TItre 1 - SMMAR']
			paraTitre1.add_run(str(oPro))
			documentStart = paraTitre1

			## Classes
			paraTitre2 = insert_paragraph_after(documentStart)
			paraTitre2.style = document.styles['TItre 2 SMMAR']
			paraTitre2.add_run('Classes')
			documentStart = paraTitre2
			for oCls in oPro.tclassesecoleprojet_set.all():
				paraNormal = insert_paragraph_after(documentStart)
				paraNormal.style = document.styles['Corps texte - SMMAR']
				paraNormal.add_run(' - '.join([
					str(oCls.get_classe()),
					str(oCls.get_ecole())
				]))	
				documentStart = paraNormal

			## Animations
			paraTitre2 = insert_paragraph_after(documentStart)
			paraTitre2.style = document.styles['TItre 2 SMMAR']
			paraTitre2.add_run('Animations')
			documentStart = paraTitre2

			### Animation
			for oAni in qsAni.filter(id_projet=oPro.pk):

				# Titre de l'animation
				aniTitre = ' '.join([
					'Le', oAni.get_dt_anim__str(), str(oAni)
				])
				oBil = None
				if oAni.get_bilan__object():
					oBil = oAni.get_bilan__object().get_ba()
				if oBil:
					aniTitre += ' - {}'.format(oBil.get_titre_ba())
				paraTitre3 = insert_paragraph_after(documentStart)
				paraTitre3.style = document.styles['Titre 3 - SMMAR']
				paraTitre3.add_run(aniTitre)
				documentStart = paraTitre3

				if oBil:

					# Nombre de personnes présentes
					if oBil.get_nbre_pers_pres_ba() is not None:
						paraNormal = insert_paragraph_after(documentStart)
						paraNormal.style = document.styles['Corps texte - SMMAR']
						run = paraNormal.add_run('Nombre de personnes présentes à l\'animation :')
						run.bold = True
						paraNormal.add_run(' {}'.format(oBil.get_nbre_pers_pres_ba()))
						documentStart = paraNormal

					# Thématiques abordées
					paraNormal = insert_paragraph_after(documentStart)
					paraNormal.style = document.styles['Corps texte - SMMAR']
					run = paraNormal.add_run('Thématiques abordées')
					run.bold = True
					documentStart = paraNormal
					paraNormal = insert_paragraph_after(documentStart)
					paraNormal.style = document.styles['Corps texte - SMMAR']
					paraNormal.add_run(oBil.get_themat_abord_ba())
					paraNormal.alignment = WD_ALIGN_PARAGRAPH.LEFT
					documentStart = paraNormal

					# Déroulement et méthodes adoptées
					if oBil.get_deroul_ba():
						paraNormal = insert_paragraph_after(documentStart)
						paraNormal.style = document.styles['Corps texte - SMMAR']
						run = paraNormal.add_run('Déroulement et méthodes adoptées')
						run.bold = True
						documentStart = paraNormal
						paraNormal = insert_paragraph_after(documentStart)
						paraNormal.style = document.styles['Corps texte - SMMAR']
						paraNormal.add_run(oBil.get_deroul_ba())
						paraNormal.alignment = WD_ALIGN_PARAGRAPH.LEFT
						documentStart = paraNormal

					# Y a-t-il eu des activités en intérieur ?
					paraNormal = insert_paragraph_after(documentStart)
					paraNormal.style = document.styles['Corps texte - SMMAR']
					run = paraNormal.add_run('Y a-t-il eu des activités en intérieur ?')
					run.bold = True
					paraNormal.add_run(' {}'.format('Oui' if oBil.get_en_inter() else 'Non'))
					documentStart = paraNormal

					# Y a-t-il eu des activités en extérieur ?
					paraNormal = insert_paragraph_after(documentStart)
					paraNormal.style = document.styles['Corps texte - SMMAR']
					run = paraNormal.add_run('Y a-t-il eu des activités en extérieur ?')
					run.bold = True
					paraNormal.add_run(' {}'.format('Oui' if oBil.get_en_exter() else 'Non'))
					documentStart = paraNormal

					# Points positifs de l'animation
					qsPP = oBil.get_point().exclude(comm_pos_point = '')
					if qsPP.exists():
						paraNormal = insert_paragraph_after(documentStart)
						paraNormal.style = document.styles['Corps texte - SMMAR']
						run = paraNormal.add_run('Points positifs de l\'animation')
						run.bold = True
						documentStart = paraNormal
						for oPP in qsPP:
							paraNormal = insert_paragraph_after(documentStart)
							paraNormal.style = document.styles['Corps texte - SMMAR']
							paraNormal.add_run(
								'{} : {}'.format(
									oPP.get_int_point(),
									oPP.get_comm_pos_point()
								)
							)
							documentStart = paraNormal

					# Points négatifs de l'animation
					'''
					qsPN = oBil.get_point().exclude(comm_neg_point = '')
					if qsPN.exists():
						paraNormal = insert_paragraph_after(documentStart)
						paraNormal.style = document.styles['Corps texte - SMMAR']
						run = paraNormal.add_run('Points négatifs de l\'animation')
						run.bold = True
						documentStart = paraNormal
						for oPN in qsPN:
							paraNormal = insert_paragraph_after(documentStart)
							paraNormal.style = document.styles['Corps texte - SMMAR']
							paraNormal.add_run(
								'{} : {}'.format(
									oPN.get_int_point(),
									oPN.get_comm_neg_point()
								)
							)
							documentStart = paraNormal
					'''

					# Images
					photos = []
					if oBil.get_photo_1_ba():
						photos.append(oBil.get_photo_1_ba())
					if oBil.get_photo_2_ba():
						photos.append(oBil.get_photo_2_ba())
					if oBil.get_photo_3_ba():
						photos.append(oBil.get_photo_3_ba())
					para = insert_paragraph_after(documentStart)
					run = para.add_run()
					for photo in photos:
						try:
							run.add_picture(photo, Mm(60))
							run.add_text(' ')
						except:
							pass
					para.alignment = WD_ALIGN_PARAGRAPH.CENTER
					documentStart = para

		# Suppression du paragraphe $START$
		p = paragraphToRemove._element
		p.getparent().remove(p)
		p._p = p._element = None

		# Préparation du téléchargement du bilan
		output = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		output['Content-Disposition'] = 'attachment; filename={}.docx'.format(gen_string())
		document.save(output)

		return output

	def __get_data(self):

		"""Récupération des données du tableau"""

		# Imports
		from django.urls import reverse

		# Initialisation des données
		data = []

		# Animations
		qsAni = self.__get_queryset()

		# Pour chaque enregistrement...
		for oAni in qsAni:

			# Définition des données
			_data = {
				'_link': '''
				<a
					href="{}"
					class="inform-icon pull-right"
					target="_blank"
					title="Consulter l'animation"
				></a>
				'''.format(
					reverse('consult_anim', args=[oAni.pk])
				),
				'aniPresta': oAni.get_projet().get_org(),
				'aniMarche': oAni.get_projet().get_pm().get_marche(),
				'marcheLot': oAni.get_projet().get_pm().get_prests(),
				'aniProjet': oAni.get_projet(),
				'aniLieu': oAni.get_lieu_anim(),
				'aniCommu': oAni.get_comm(),
				'aniDate': oAni.get_dt_anim__str()
			}

			# Empilement des données
			data.append(_data)

		return data

	def __get_datatable(self):

		"""Table HTML"""

		# Imports

		# Données filtrées
		data = self.__get_data()

		# Balise </tbody>
		trs = []
		for i in data:
			_tds = [
				i['aniPresta'],
				i['aniMarche'],
				i['marcheLot'],
				i['aniProjet'],
				i['aniLieu'],
				i['aniCommu'],
				i['aniDate'],
				i['_link']
			]
			tds = ''.join(['<td>{}</td>'.format(j) for j in _tds])
			tr = '<tr>{}</tr>'.format(tds)
			trs.append(tr)
		tbody = ''.join(trs)

		return '''
		<div
			class="custom-table"
			id="dtable_real_etats_Bilan_Animations_Facturation"
		>
			<table border="1" bordercolor="#DDD">
				<thead>
					<tr>
						<th>Organisme en charge de l'animation</th>
						<th>Marché</th>
						<th>Lot</th>
						<th>Projet</th>
						<th>Lieu de l'animation</th>
						<th>Commune accueillant l'animation</th>
						<th>Date de l'animation</th>
						<th></th>
					</tr>
				</thead>
				<tbody>{}</tbody>
			</table>
		</div>
		'''.format(tbody)

	def __get_form(self):

		"""Formulaire"""

		# Imports
		from app.functions.form_init import sub as form_init
		from django.template.context_processors import csrf

		# Initialisation des contrôles
		form = form_init(self)

		return '''
		<form
			action=""
			method="post"
			name="form_real_etats_Bilan_Animations_Facturation"
			onsubmit="ajax(event);"
		>
			<input name="csrfmiddlewaretoken" type="hidden" value="{}">
			<fieldset class="my-fieldset">
				<legend>Filtrer par</legend>
				<div>
					{}
					{}
					{}
					{}
					<button
						class="center-block custom-button main-button"
						type="submit"
					>Valider</button>
				</div>
			</fieldset>
		</form>
		'''.format(
			csrf(self.rq)['csrf_token'],
			form['zl_presta'],
			form['zl_marche_lot'],
			form['zd_date_debut'],
			form['zd_date_fin']
		)

	def __get_queryset(self):

		"""Jeu de données"""

		# Imports
		from app.models import TAnimation
		from django.db.models import Q


		# Initialisation des données
		data = []

		# Requête HTTP "POST" ou requête HTTP "GET" édition du bilan
		if self.data or self.rq.GET.get('action') == 'editer-bilan':

			# Récupération des données nettoyées du formulaire
			cleaned_data = self.__cleaned_data()

			# Filtres
			ands = {}
			ands['id_projet__id_org'] = cleaned_data['zl_presta']
			ands['id_projet__id_pm'] = cleaned_data['zl_marche_lot']
			ands['dt_anim__gte'] = cleaned_data['zd_date_debut']
			ands['dt_anim__lte'] = cleaned_data['zd_date_fin']

			# Jeu de données filtré
			qsAni = TAnimation.objects.filter(**ands)

		# Requête HTTP "GET" affichage vue
		else:

			# Jeu de données vierge
			qsAni = TAnimation.objects.none()

		return qsAni

	# Méthodes publiques

	def edit_bilan(self):
		"""Bilan des animations au format Word"""
		return self.__edit_bilan()

	def get_datatable(self):
		"""Table HTML"""
		return self.__get_datatable()

	def get_form(self):
		"""Formulaire"""
		return self.__get_form()